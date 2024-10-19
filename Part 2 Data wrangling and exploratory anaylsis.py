import pandas as pd #for handling tables
import numpy as np
import warnings
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from docx import Document
from docx.shared import Inches
import os

# Suppress all warnings
warnings.filterwarnings("ignore")


#Data collection
df = pd.read_csv("https://cf-courses-data.s3.us.cloud-object-storage.appdomain.cloud/IBM-DA0321EN-SkillsNetwork/LargeData/m2_survey_data.csv")
#print("\n\n",df.head())
#print(df[["Respondent", "Age", "Country"]].describe(include="all"))
#print("\n\n",df.info())



# Define the file name in which to append the results/output
doc_file = "Part_5_Presentation_of_results.docx"
#retireve program name
program_name = os.path.splitext(os.path.basename(__file__))[0]

# Function to update results into a document 
def create_or_update_doc(results_texts, graphs):
    # Check if the document exists
    if os.path.exists(doc_file):
        document = Document(doc_file)
    else:
        print(f"Document '{doc_file}' not found. Please ensure the document exists.")
        return  # Exit the function if the document is not found
    document.add_heading("\n\n\n" + program_name, level=2)
    # Add the text results
    for i, results_text in enumerate(results_texts, start=1):
        results_section_title = f"Data Analysis Log {i}"
        # Add new section and text
        document.add_heading(results_section_title, level=4)
        document.add_paragraph(results_text)

    # Add graphs (save them to disk manually)
    for i in plt.get_fignums():  # Get all figure numbers
        fig = plt.figure(i)
        
        # Save the figure as a temporary image file
        graph_path = f'temp_fig_{i}.png'
        fig.savefig(graph_path, dpi=100)  # Save with a reasonable dpi for better quality

        # Add new graph section and image
        graph_section_title = f"Graph {i+1}"
        document.add_heading(graph_section_title, level=4)
        document.add_paragraph(f"Description of Graph {i+1}")
        document.add_picture(graph_path, width=Inches(5.5))  # Add the image from the saved file

        # Optionally remove the image file after inserting it into the document
        os.remove(graph_path)

        # Close the figure to free up memory
        plt.close(fig)

    # Save the document
    document.save(doc_file)

graphs=[]













#Data Cleaning/wrangling
#to keep in mind how big the data is after every action affecting its size
num_rows = len(df)
#print(f"\n\nThe DataFrame has {num_rows} rows at the start.\n\n")

#handling duplicates
duplicates=df.duplicated()
num_duplicates = duplicates.sum() #counts only the 'True' values (duplicates)
#print(f"\n\nThere are {num_duplicates} duplicates")
if num_duplicates>0:
    #df.drop_duplicates(inplace=True)
    # Recalculate duplicates to check if any are still there
    new_duplicates = df.duplicated().sum()
    if new_duplicates == 0:
        print("\n\nDuplicates deleted\n\n")
    else:
        print("\n\nDuplicates still there")


#to keep in mind how big the data is after every action affecting its size
num_rows = len(df)
#print(f"\n\nThe DataFrame has {num_rows} rows after deleting duplicates.\n\n")



#handling missing values
missing=df.isnull()
#handling specific columns
overview=df[["Country","EdLevel"]].isnull().sum()
print(overview)
workloc_missing= missing["WorkLoc"].sum()
if workloc_missing>0:
    print(f"\n\nThere were {workloc_missing} missing values in the Workloc column")
    most_frequent_workloc = df["WorkLoc"].mode()[0]
    df["WorkLoc"]=df["WorkLoc"].replace(np.nan,most_frequent_workloc)
    df["WorkLoc"].replace(np.nan,most_frequent_workloc,inplace=True)
    new_workloc_missing= df["WorkLoc"].isnull().sum()
    if new_workloc_missing==0:
        print("\n\nrepaired workloc column")
    else:
        print("\n\nworkloc still got missing values")


num_missing = missing.sum().sum() #counts missing values in each coulmn which then gets add up for the entrie df
print(f"\n\nThere are {num_missing} missing values")
if num_missing>0:
    #df.dropna(inplace=True)
    # Recalculate duplicates to check if any are still there
    new_missing = df.isnull().sum().sum()
if new_missing== 0:
    print("\n\nALL missing deleted")
else:
    print("\n\nmissing still there")



#to keep in mind how big the data is after every action affecting its size
num_rows = len(df)
#print(f"\n\nThe DataFrame has {num_rows} rows after deleting missing rows.\n\n")




#normalizing data: adding a uniform Compensation column
categories_CompFreq = df["CompFreq"].unique()
#print("\n\n",f"There are {len(categories_CompFreq)} diffrent compensation categories.These include {categories_CompFreq}","\n\n")

# Initialize the new column with zeros or NaNs to avoid missing data issues
df["NormalizedAnnualCompensation"] = pd.NA
#access .loc[rows,column] where rows von compfreq is e.g monthly and set the value of
# Normalized compensation column to value of Comptotal*12
df.loc[df["CompFreq"] == "Yearly", "NormalizedAnnualCompensation"] = df["CompTotal"]
df.loc[df["CompFreq"] == "Monthly", "NormalizedAnnualCompensation"] = df["CompTotal"] * 12
df.loc[df["CompFreq"] == "Weekly", "NormalizedAnnualCompensation"] = df["CompTotal"] * 52



















#Exploratory Data anaylsis


#distribution of data

yearly_salaries = df[["Age", "ConvertedComp"]].dropna() 
yearly_salaries.sort_values(by="Age", inplace=True)

# Group by Age and Salary to count the number of responders for each combination
grouped_salaries = yearly_salaries.groupby(["Age", "ConvertedComp"]).size().reset_index(name='RespondentCount')

fig, (ax0, ax1) = plt.subplots(1, 2, figsize=(15, 6))

# Scatterplot: Bubble size reflects number of respondents with the same salary
bubble_size = grouped_salaries['RespondentCount'] * 100  # Adjust the scaling of the bubble size
sns.scatterplot(data=grouped_salaries, x='Age', y='ConvertedComp', ax=ax0, size=bubble_size, 
                sizes=(50, 1500), alpha=0.7, legend=False)

# Scale the y-axis to hundreds of thousands instead of millions
ax0.set_yticklabels(['{:.0f}e5'.format(y/1e5) for y in ax0.get_yticks()])
ax0.set_xlabel('Age of Respondents', fontsize=12)
ax0.set_ylabel('Yearly Salaries (in hundreds of thousands)', fontsize=12)
ax0.set_title('Distribution of Annual Salaries by Age', fontsize=14)
ax0.annotate('Bubble size reflects the number of respondents with the same salary', 
             xy=(0.5, -0.12), xycoords='axes fraction', ha='center', fontsize=10, color='gray')

# Histogram: Salary distribution with aligned bins and ticks
num_bins = 10
counts, bin_edges, _ = ax1.hist(yearly_salaries['ConvertedComp'], bins=num_bins, edgecolor='black', color='skyblue')

ax1.set_xticks(bin_edges)  # Align ticks with bin edges
ax1.set_xticklabels(['{:.0f}e5'.format(x/1e5) for x in bin_edges], rotation=45)
ax1.set_xlabel('Yearly Salaries (in hundreds of thousands)', fontsize=12)
ax1.set_ylabel('Number of Respondents', fontsize=12)
ax1.set_title('Distribution of Annual Salaries', fontsize=14)

# Add grid for better readability
ax0.grid(True, linestyle='--', alpha=0.6)
ax1.grid(True, linestyle='--', alpha=0.6)

plt.tight_layout()

#saving/collecting graphs
plt.savefig('scatterplot_histogram_canvas.png', dpi=100)
graphs.append('scatterplot_histogram_canvas.png')

#plt.show()

results_text1 = f"The median age is {df['Age'].median()} years old!"
results_text2=f"The median annual salary is {df["ConvertedComp"].median()} dollars!"
results_text3=f"There were {len(df.loc[df["Gender"]=="Man"])} Men!"
results_text4=f"The female median annual salary is {df.loc[df["Gender"]=="Woman"]["ConvertedComp"].median()} dollars!"

age_summary = df['Age'].describe()
results_text5 = "age summary"+age_summary.to_string()  # Convert Series to string
salary_summary = df['ConvertedComp'].describe()
results_text6 = "salary summary"+salary_summary.to_string()  # Convert Series to string




#finding and handling outliers with a boxplot
fig=plt.figure(figsize=(12, 6))
boxplot = sns.boxplot(x="Age", data=yearly_salaries, color="lightblue")
#boxplot.set_yscale('log') #logarithmic scale to the y-axis for better visualization
plt.grid(True, which="both", axis="y", linestyle="--", linewidth=0.7)
plt.title("Age of Respondents", fontsize=16)
plt.ylabel("Age (Log Scale)", fontsize=12)
plt.xlabel("Age", fontsize=12)
plt.xticks(np.arange(0, df["Age"].max(), step=5))
plt.grid(True)
#saving/collecting graphs
plt.savefig('boxplot.png', dpi=100)
graphs.append('boxplot.png')
#plt.show()

IQR=df['ConvertedComp'].quantile(0.75)-df['ConvertedComp'].quantile(0.25)
results_text7=f"Half of Responders make between {df['ConvertedComp'].quantile(0.75)} and {df['ConvertedComp'].quantile(0.25)} dollars yearly."
results_text8=f"Special cases(outliers) are People who make over {df['ConvertedComp'].quantile(0.75)+IQR*1.5} and under {df['ConvertedComp'].quantile(0.25)-IQR*1.5} dollars yearly "
df = df.loc[
    (df["ConvertedComp"] < df['ConvertedComp'].quantile(0.75) + IQR * 1.5) &
    (df["ConvertedComp"] > df['ConvertedComp'].quantile(0.25) - IQR * 1.5)
]

new_salary_summary = df['ConvertedComp'].describe()
results_text9="Salaries data considered will disregard outliers:" + str(new_salary_summary)



results_text10 = f"The realistic median annual salary is {df['ConvertedComp'].median()} dollars!"
results_text11=f"The realistic average annual salary is {df["ConvertedComp"].mean()} dollars!"


#correlation
# .corr() computes pairwise correlation of columns and returns features correlated to age
# as a correlation matrix. This can help determine which other factors closely relate to
#the age of the respondants

# Select only numeric columns for both correlation and p-value calculations
df = df.select_dtypes(include=[float, int])
# Calculate the correlation matrix inbetween all numeric columns and "Age" sorted descending
correlation_matrix = df.corr()["Age"].sort_values()
# Filter correlations to include only those greater than 
high_correlations = correlation_matrix[correlation_matrix > 0.3]
# Convert the correlation matrix to a DataFrame(only for better output format)
high_correlations_df = high_correlations.to_frame(name='Correlation')
# Print the columns with high correlation to Age
results_text12 = "Age-determined/correlated factors:\n" + high_correlations_df.to_string()  # Convert DataFrame to string


# Check p-values of the correlations to evaluate certainty
p_values = []
results_text13 = "Correlation and P-values for Age:\n"
# Loop over the columns in the high correlations and calculate their p-values
for col in high_correlations.index:
        _, p_value = stats.pearsonr(df[col], df["Age"])
        p_values.append((col, p_value))
# Print the columns and their corresponding p-values
for col, p_value in p_values:
    results_text13 += f"Column: {col}, P-value: {p_value}\n"  # Append each result to results_text



results_text14 = correlation_matrix.to_string()







# List of graphs (paths to image files) and texts
results_texts = [results_text1, results_text2,results_text3,results_text4,str(results_text5),
                 str(results_text6),results_text7,results_text8,str(results_text9),results_text10,
                 results_text11,results_text12,results_text13,results_text14]




# Call the function to create/update the document
#create_or_update_doc(results_texts,graphs)  

print("Document has been updated successfully.")
