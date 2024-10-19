import pandas as pd
import numpy as np
import warnings
import matplotlib.pyplot as plt
import seaborn as sns
import sqlite3
import os


# Set the working directory
os.chdir("C:/Users/kalab/OneDrive/Desktop/py4e/python first programms/dataanalysis IBM/capstone/.gitignore")
print("Current Working Directory:", os.getcwd())  # Check if the working directory is correct

# Suppress all warnings
warnings.filterwarnings("ignore")





# Define the file name
doc_file = "Part_5_Presentation_of_results.docx"
#retireve program name
program_name = os.path.splitext(os.path.basename(__file__))[0]

# Function to update results into a document 
from docx import Document
from docx.shared import Inches
def create_or_update_doc(results_texts, graphs): 
    # Check if the document exists
    if os.path.exists(doc_file):
        document = Document(doc_file)
    else:
        print(f"Document '{doc_file}' not found. Please ensure the document exists.")
        return  # Exit the function if the document is not found
    
    # Add the program name as a header
    document.add_heading("\n\n\n" + program_name, level=2)
    
    # Add the text results
    for i, results_text in enumerate(results_texts, start=1):
        results_section_title = f"Data Analysis Log {i}"
        document.add_heading(results_section_title, level=4)
        document.add_paragraph(results_text)

    # Add graphs (save them to disk manually)
    for graph_path in graphs:
        # Check if the graph file exists
        if os.path.exists(graph_path):
            print(f"Adding {graph_path} to the document.")  # Debugging statement
            graph_section_title = f"Graph {graphs.index(graph_path) + 1}"
            document.add_heading(graph_section_title, level=4)
            document.add_paragraph(f"Description of Graph {graphs.index(graph_path) + 1}")
            document.add_picture(graph_path, width=Inches(5.5))  # Add the image from the saved file
        else:
            print(f"File {graph_path} not found, skipping.")

    # Save the document after all images and texts are added
    try:
        document.save(doc_file)
        print(f"Document '{doc_file}' has been updated and saved.")
    except Exception as e:
        print(f"Error saving the document: {e}")
    
    # Optionally remove the image files after inserting them into the document
    for graph_path in graphs:
        if os.path.exists(graph_path):
            os.remove(graph_path)
            print(f"Temporary file {graph_path} deleted.")

graphs=[]





conn = sqlite3.connect("m4_survey_data.sqlite") # open a database connection
cur = conn.cursor()





#Getting an overview of a database using panda

# -- Query 1: To find the number of tables and their names
#sqlite_master is a system table that stores metadata

#tables_query = "SELECT name FROM sqlite_master WHERE type='table';"
#tables = pd.read_sql_query(tables_query, conn)

# Check if 'tables' DataFrame is created and contains the table names
#if not tables.empty:
    #print(f"Tables in the database: {tables['name'].tolist()}")

    # Iterate over each table to get the number of rows, columns, and column names
    #for table in tables['name']:
        #print(f"\nTable: {table}")
        
        # Get the number of columns and column names
        #columns_info = pd.read_sql_query(f"PRAGMA table_info({table});", conn)
        #num_columns = len(columns_info)
        #column_names = columns_info['name'].tolist()
        
        # Get the number of rows
        #num_rows = pd.read_sql_query(f"SELECT COUNT(*) FROM {table};", conn).iloc[0, 0]
        
        # Output the information
        #print(f"\n\nNumber of Columns: {num_columns}")
        #print(f"\n\nColumn Names: {column_names}")
        #print(f"\n\nNumber of Rows: {num_rows}\n\n")
#else:
#print("No tables found in the database.")

# Close the database connection









#visualizing distribution and composition of data
salary_q = "Select CompTotal from master;"  # Ensure table name is correct
salary_df = pd.read_sql_query(salary_q, conn)
#results_text1="Salary Overview"+salary_df.head().to_string()


salary_df = salary_df.dropna()
salary_df.sort_values(by="CompTotal", inplace=True)

# Plot the histogram with a log scale on the x-axis
plt.figure(figsize=(8, 6))
plt.hist(salary_df['CompTotal'], bins=50, edgecolor='black')
plt.xlabel('Yearly salaries (in hundreds of thousands)')
plt.ylabel('Number of responders')
plt.title('Distribution of Annual Salaries')
tick_locations = [0, 1e5, 2e5, 3e5, 4e5, 5e5]  # Replace with the actual range of your data
tick_labels = ['0', '100k', '200k', '300k', '400k', '500k']
plt.xticks(tick_locations, tick_labels)
#plt.show()
plt.savefig('Salary_distribution_respondents.png', dpi=100)
graphs.append('Salary_distribution_respondents.png')
plt.close() # figure remains open in memory when using plt.savefig() until closed, which can consume resources




#finding and handling outliers with a boxplot
age_q = "Select Age from master;"  # Ensure table name is correct
age_df = pd.read_sql_query(age_q, conn)
#print(age_df.head())

plt.figure(figsize=(8, 6))  # Adjust the figure size for better visualization
boxplot = sns.boxplot(x=age_df["Age"], color="lightblue")
plt.grid(True, which="both", axis="y", linestyle="--", linewidth=0.7)
plt.title("Distribution of Respondents' Age", fontsize=16, fontweight='bold', pad=20)
plt.ylabel("Number of Respondents", fontsize=12)
plt.xlabel("Age of Respondents", fontsize=12)
#plt.show()
plt.savefig('Age_distribution_respondents.png', dpi=100)
graphs.append('Age_distribution_respondents.png')
plt.close()


#most popular languages respondents wish to learn next year,
lang_q="""SELECT LanguageDesireNextYear, COUNT(LanguageDesireNextYear) AS count 
FROM LanguageDesireNextYear 
GROUP BY LanguageDesireNextYear
ORDER BY count DESC
LIMIT 5;"""
lang_df=pd.read_sql_query(lang_q,conn)

plt.figure(figsize=(8, 6))
labels = lang_df["LanguageDesireNextYear"]
sizes = lang_df["count"]
plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
plt.title('Languages Respondents Wish to Learn', fontsize=14)
#plt.show()
plt.savefig('languages_respondents_wish_to_learn.png', dpi=100)
graphs.append('languages_respondents_wish_to_learn.png')
plt.close()








#describing Relationships

#relationship between age and WorkWeekHrs
r_q1="SELECT Age, WorkWeekHrs FROM master;"
r_df = pd.read_sql_query(r_q1, conn)
#print(r_df.head())

fig = plt.figure(figsize=(12, 6))
size = r_df['WorkWeekHrs'] * 0.2  # Adjust the scaling of the bubble size for better visibility
plt.scatter(r_df['Age'], r_df['WorkWeekHrs'], s=size, color="darkblue", alpha=0.8)
plt.xlabel('Age of Respondents', fontsize=14)
plt.ylabel('Work Week Hours', fontsize=14)
plt.xlim(right=80)
plt.title('Distribution of Work Week Hours by Age', fontsize=16)
plt.grid(True)
#plt.show()
plt.savefig('relationship between age and WorkWeekHrs.png', dpi=100)
graphs.append('relationship between age and WorkWeekHrs.png')
plt.close()



#databases respondents wish to learn
r_q2="""SELECT COUNT(DatabaseDesireNextYear) as count, DatabaseDesireNextYear 
FROM DatabaseDesireNextYear
GROUP BY DatabaseDesireNextYear 
ORDER BY count DESC 
LIMIT 5;"""
r_df2 = pd.read_sql_query(r_q2, conn)
#print(r_df2.head())

plt.figure(figsize=(8, 6))
labels = r_df2["DatabaseDesireNextYear"]
sizes = r_df2["count"]
plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
plt.title('databases respondents wish to learn')
#plt.show()
plt.savefig('databases_respondents_wish_to_learn.png', dpi=100)
graphs.append('databases_respondents_wish_to_learn.png')
plt.close()


#work and coderevhrs
r_q3="SELECT Age, WorkWeekHrs, CodeRevHrs from master where Age BETWEEN 30 and 35;"
r_df3 = pd.read_sql_query(r_q3, conn)
#print(r_df3.head())

median_by_age = r_df3.groupby('Age')[['WorkWeekHrs', 'CodeRevHrs']].median().reset_index()

median_by_age.set_index('Age').plot(kind='bar', figsize=(10, 6))
plt.xlabel('Age')
plt.ylabel('Median Hours')
max_hours=median_by_age["WorkWeekHrs"].max() + 15
plt.ylim(0,max_hours)
plt.yticks(np.arange(0, max_hours , step=5))
plt.title('Median Work Week Hours vs Code Review Hours (Aged 30-35)')
#plt.show()
plt.savefig('Median_work_and_code-review_hours_per_week.png', dpi=100)
graphs.append('Median_work_and_code-review_hours_per_week.png')
plt.close()









#visualizing comparison of data
r_q4 = """
SELECT Age, MainBranch, ConvertedComp
FROM master
WHERE Age BETWEEN 25 AND 60;
"""
r_df4 = pd.read_sql_query(r_q4, conn)

# Calculate the median ConvertedComp for each age between 25 and 60
median_convertedcomp = r_df4.groupby("Age")["ConvertedComp"].median().reset_index()

plt.figure(figsize=(10, 6))
plt.plot(median_convertedcomp["Age"], median_convertedcomp["ConvertedComp"], marker='o')
plt.title("Median ConvertedComp for Ages 25 to 60")
plt.xlabel("Age")
plt.ylabel("Median ConvertedComp")
plt.grid(True)
#plt.show()
plt.savefig('Median_annual_income_by_age.png', dpi=100)
graphs.append('Median_annual_income_by_age.png')
plt.close()

# Count the occurrences of each MainBranch
mainbranch_counts = r_df4["MainBranch"].value_counts().reset_index()
mainbranch_counts.columns = ["MainBranch", "Count"]

plt.figure(figsize=(10, 6))
plt.barh(mainbranch_counts["MainBranch"], mainbranch_counts["Count"], color='skyblue')
plt.title("MainBranch Distribution for Ages 45 to 60", fontsize=14)
plt.xlabel("Count", fontsize=12)
plt.ylabel("MainBranch", fontsize=12)
plt.subplots_adjust(left=0.3)  # Increase left margin to allow space for long y-axis labels
plt.grid(True, axis='x')
#plt.show()
plt.savefig('fulltime_vs_parttime_coders.png', dpi=100)
graphs.append('fulltime_vs_parttime_coders.png')
plt.close()


# Majority of the survey responders are
r_q4 = "SELECT DevType, count(DevType) as count from DevType group by DevType order by count desc;"
r_df4 = pd.read_sql_query(r_q4, conn)
#print(r_df4.head())


plt.figure(figsize=(12, 8))
y = r_df4["count"]  # Counts for x-axis
x = range(len(r_df4))  # Numerical indices for y-axis, as labels cannot be used directly in scatter plots
plt.scatter(x, y, marker='o')

plt.xticks(ticks=range(len(r_df4["DevType"])), labels=r_df4["DevType"], rotation=55, ha='right')
plt.title('Type of Survey Responders',fontsize=14)
plt.ylabel("Amount",fontsize=12)
plt.xlabel("Dev Type",fontsize=12)
plt.tight_layout(pad=2.0)
plt.show()
plt.savefig('resoponder_roles_scatter.png', dpi=100)
graphs.append('resoponder_roles_scatter.png')
plt.close() 


# List of graphs (paths to image files) and texts
results_texts = []




# Call the function to create/update the document

#Debugging control points
for graph in graphs:
    if os.path.exists(graph):
        print(f"File {graph} exists.")
    else:
        print(f"File {graph} not found.")
print("Graphs list:", graphs)  


#create_or_update_doc(results_texts,graphs)  

print("\n\nDocument has been updated successfully.")




conn.close()