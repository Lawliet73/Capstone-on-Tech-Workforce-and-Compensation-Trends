from docx import Document
from docx.shared import Inches
from docx.shared import Inches
import pandas as pd #for handling tables
import numpy as np
import matplotlib.pyplot as plt
import warnings
import os

# Suppress all warnings
warnings.filterwarnings("ignore")

url = "https://cf-courses-data.s3.us.cloud-object-storage.appdomain.cloud/IBM-DA0321EN-SkillsNetwork/LargeData/m1_survey_data.csv"

df=pd.read_csv(url)


# Define the file name
doc_file = "Part_5_Presentation_of_results.docx"

# Function to update results into a document
def create_or_update_doc(results_text, graphs):
    # Check if the document exists, otherwise create it
    if os.path.exists(doc_file):
        document = Document(doc_file)
    else:
        document = Document()

    # Function to check if a section already exists
    def section_exists(document, section_title):
        for paragraph in document.paragraphs:
            if paragraph.text == section_title:
                return True
        return False

    # Function to replace a section(text and image) if it already exists
    def replace_section(document, section_title, new_text):
        found = False
        for i, paragraph in enumerate(document.paragraphs):
            if paragraph.text == section_title:
                found = True
                index = i
                break
        
        if found:
            # Remove everything (text and images) under this section
            while index + 1 < len(document.paragraphs) and document.paragraphs[index + 1].style.name != 'Heading 1':
                del document.paragraphs[index + 1]
            # Insert new text after the title
            document.paragraphs[index].insert_paragraph_after(new_text)

    # Add or replace multiple text result sections marked as results_textn
    for i, results_text in enumerate(results_texts, start=1):
        results_section_title = f"Data Processing and Modifications Log {i}"
        if section_exists(document, results_section_title):
            replace_section(document, results_section_title, results_text)
        else:
            document.add_heading(results_section_title, level=4)
            document.add_paragraph(results_text)

     # Automatically collect and add all open matplotlib figures into the docx
    for i in plt.get_fignums():  # Get all figure numbers
        fig = plt.figure(i)
        graph_path = f'temp_fig_{i}.png'  # Temporary image file for the graph
        fig.savefig(graph_path)  # Save the figure as an image file

        graph_section_title = f"Graph {i+1}"
        if section_exists(document, graph_section_title):
            # Replace existing section and image
            replace_section(document, graph_section_title, f"Updated Graph {i+1}")
        else:
            # Add new section and image
            document.add_heading(graph_section_title, level=2)
            document.add_paragraph(f"Description of Graph {i+1}")
            document.add_picture(graph_path, width=Inches(5.5))

        # Remove temporary image file after inserting into the document
        os.remove(graph_path)

    # Save the document
    document.save(doc_file)



#print("\n\n",df.head())
#print(df[["Respondent", "Age", "Country"]].describe(include="all"))
#print("\n\n",df.info())



# Get all column names
column_names = df.columns.tolist()
column_string = ', '.join(column_names)
result_text1 = f"The analyzed data has information on {column_string}."


#to keep in mind how big the data is after every action affecting its size
num_rows = len(df)
result_text2=f"\n\nThe DataFrame has {num_rows} rows at the start.\n\n"

#handling duplicates
duplicates=df.duplicated()
num_duplicates = duplicates.sum() #counts only the 'True' values (duplicates)
result_text3=f"\n\nThere are {num_duplicates} duplicates"
if num_duplicates>0:
    df.drop_duplicates(inplace=True)
    # Recalculate duplicates to check if any are still there
    new_duplicates = df.duplicated().sum()
if new_duplicates== 0:
    result_text4="\n\nAll Duplicates were deleted.\n\n"
else:
    print("\n\nDuplicates still there")


#to keep in mind how big the data is after every action affecting its size
num_rows = len(df)
result_text5=f"\n\nThe DataFrame has {num_rows} rows after deleting duplicates.\n\n"



#handling missing values
missing=df.isnull()
overview=df[["Country","EdLevel"]].isnull().sum()
#print(overview)
workloc_missing= missing["WorkLoc"].sum()
if workloc_missing>0:
    #print(f"\n\nThere were {workloc_missing} missing values in the Workloc column")
    most_frequent_workloc = df["WorkLoc"].mode()[0]
    df["WorkLoc"]=df["WorkLoc"].replace(np.nan,most_frequent_workloc)
    df["WorkLoc"].replace(np.nan,most_frequent_workloc,inplace=True)
    new_workloc_missing= df["WorkLoc"].isnull().sum()
    if new_workloc_missing==0:
        result_text6="\n\nrepaired workloc column"
    else:
        print("\n\nworkloc still got missing values")


num_missing = missing.sum().sum() #counts missing values in each coulmn which then gets add up for the entrie df
result_text7=f"\n\nThere are {num_missing} missing values in columns which can´t be replaced"
if num_missing>0:
    df.dropna(inplace=True)
    # Recalculate duplicates to check if any are still there
    new_missing = df.isnull().sum().sum()
if new_missing== 0:
    result_text8="\n\nALL rows with missing values deleted"
else:
    print("\n\nmissing still there")



#to keep in mind how big the data is after every action affecting its size
num_rows = len(df)
result_text9=f"\n\nThe DataFrame has {num_rows} rows after deleting missing rows.\n\n"


#normalinzing data
categories_CompFreq = df["CompFreq"].unique()
result_text10 = f"\n\nThere are {len(categories_CompFreq)} different compensation categories. These include {categories_CompFreq}"

# Initialize the new column with zeros or NaNs to avoid missing data issues
df["NormalizedAnnualCompensation"] = pd.NA
#access .loc[rows,column] where rows von compfreq is e.g monthly and set the value of
# Normalized compensation column to value of Comptotal*12
df.loc[df["CompFreq"] == "Yearly", "NormalizedAnnualCompensation"] = df["CompTotal"]
df.loc[df["CompFreq"] == "Monthly", "NormalizedAnnualCompensation"] = df["CompTotal"] * 12
df.loc[df["CompFreq"] == "Weekly", "NormalizedAnnualCompensation"] = df["CompTotal"] * 52
result_text11="A normalized Annual Compensation cloumn was created to reflect a comparable income of responders"

#print("\n\n",df["NormalizedAnnualCompensation"].head(),"\n\n")




# List of graphs (paths to image files) and texts
results_texts = [result_text1, result_text2,result_text3,result_text4,result_text5,
                 result_text6,result_text7,result_text8,result_text9,result_text10,
                 result_text11]
# Call the function to create/update the document
create_or_update_doc(results_texts,[])  #no graphs so pass empty string

print("Document has been updated successfully.")







